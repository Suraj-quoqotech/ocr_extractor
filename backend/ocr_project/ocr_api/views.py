from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.permissions import IsAuthenticated
from rest_framework_simplejwt.authentication import JWTAuthentication
from rest_framework_simplejwt.views import TokenObtainPairView
from django.conf import settings
from pathlib import Path
import os
import json
from mistralai import Mistral
from dotenv import load_dotenv
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from django.contrib.auth.models import User
from django.utils import timezone

from .models import OCRFile
from django.http import FileResponse, Http404
from django.shortcuts import get_object_or_404

from rest_framework import status
from rest_framework.views import APIView
from rest_framework.response import Response
from .serializers import RegisterSerializer
from rest_framework.permissions import AllowAny

from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from django.contrib.auth.models import User
from .models import ChatRoom, Message
from .serializers import ChatUserSerializer, ChatRoomSerializer, MessageSerializer

class RegisterView(APIView):
    permission_classes = [AllowAny]
    def post(self, request):
        serializer = RegisterSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({"msg":"user created"}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

# Load API key from .env
load_dotenv()
API_KEY = os.getenv("MISTRAL_API_KEY")
if not API_KEY:
    raise ValueError("MISTRAL_API_KEY not found in .env")

client = Mistral(api_key=API_KEY)


class OCRUploadView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, format=None):
        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({"error": "No file uploaded"}, status=400)
        # Get selected output formats from request
        output_formats_json = request.POST.get('output_formats', '["pdf", "txt", "docx"]')
        try:
            output_formats = json.loads(output_formats_json)
        except:
            output_formats = ["pdf", "txt", "docx"]  # default to all
        
        # Validate at least one format is selected
        if not output_formats or len(output_formats) == 0:
            return Response({"error": "At least one output format must be selected"}, status=400)

        # Save uploaded file temporarily
        user_upload_dir = Path(settings.MEDIA_ROOT) / "uploads" / str(request.user.id)
        user_upload_dir.mkdir(parents=True, exist_ok=True)
        input_path = user_upload_dir / file_obj.name
        with open(input_path, "wb") as f:
            for chunk in file_obj.chunks():
                f.write(chunk)

        # Upload to Mistral Files API
        with open(input_path, "rb") as f:
            uploaded = client.files.upload(
                file={"file_name": file_obj.name, "content": f},
                purpose="ocr"
            )

        signed = client.files.get_signed_url(file_id=uploaded.id, expiry=1)
        url = signed.url

        # OCR processing
        ocr_response = client.ocr.process(
            model="mistral-ocr-latest",
            document={"type": "document_url", "document_url": url},
            include_image_base64=False
        )

        all_text = [page.markdown for page in ocr_response.pages]
        final_text = "\n\n".join(all_text)

        # Prepare output folder
        user_output_dir = Path(settings.MEDIA_ROOT) / "outputs" / str(request.user.id)
        user_output_dir.mkdir(parents=True, exist_ok=True)

        base_name = Path(file_obj.name).stem

        # Initialize response data
        response_data = {}
        
        # Generate PDF if requested
        pdf_url = None
        pdf_size = 0
        if "pdf" in output_formats:
            pdf_path = user_output_dir / f"ocr_{base_name}.pdf"
            doc = SimpleDocTemplate(str(pdf_path), pagesize=A4,
                                    rightMargin=40, leftMargin=40,
                                    topMargin=40, bottomMargin=40)
            story = []
            styles = getSampleStyleSheet()
            normal_style = styles['Normal']

            for paragraph in final_text.split("\n\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    story.append(Paragraph(paragraph.replace("\n", "<br/>"), normal_style))
                    story.append(Spacer(1, 12))
            doc.build(story)
            
            pdf_url = request.build_absolute_uri(settings.MEDIA_URL + f"outputs/{pdf_path.name}")
            pdf_size = pdf_path.stat().st_size
            response_data['pdf'] = pdf_url
            response_data['pdf_size'] = pdf_size

        # Generate TXT if requested
        txt_url = None
        txt_size = 0
        if "txt" in output_formats:
            txt_path = user_output_dir / f"ocr_{base_name}.txt"
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(final_text)
            
            txt_url = request.build_absolute_uri(settings.MEDIA_URL + f"outputs/{txt_path.name}")
            txt_size = txt_path.stat().st_size
            response_data['txt'] = txt_url
            response_data['txt_size'] = txt_size

        # Generate DOCX if requested
        docx_url = None
        docx_size = 0
        if "docx" in output_formats:
            docx_path = user_output_dir / f"ocr_{base_name}.docx"
            docx_doc = Document()
            for paragraph in final_text.split("\n\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    docx_doc.add_paragraph(paragraph)
                    docx_doc.add_paragraph("")
            docx_doc.save(docx_path)
            
            docx_url = request.build_absolute_uri(settings.MEDIA_URL + f"outputs/{docx_path.name}")
            docx_size = docx_path.stat().st_size
            response_data['docx'] = docx_url
            response_data['docx_size'] = docx_size
        # Save to DB and associate with requesting user
        OCRFile.objects.create(
        file_name=file_obj.name,
        pdf_url=pdf_url or "",
        txt_url=txt_url or "",
        docx_url=docx_url or "",
        pdf_size=pdf_size,
        txt_size=txt_size,
        docx_size=docx_size,
        status="done",
        user=request.user,
        )


        return Response(response_data)


class OCRHistoryView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # Admin user (username 'Admin') sees all files
        if request.user.username == 'Admin':
            files = OCRFile.objects.all().order_by('-uploaded_at')
        else:
            files = OCRFile.objects.filter(user=request.user).order_by('-uploaded_at')
        data = []

        for f in files:
            item = {
                "file_name": f.file_name,
                "pdf_url": f.pdf_url if f.pdf_url else None,
                "txt_url": f.txt_url if f.txt_url else None,
                "docx_url": f.docx_url if f.docx_url else None,
                "pdf_size": f.pdf_size,
                "txt_size": f.txt_size,
                "docx_size": f.docx_size,
                "uploaded_at": f.uploaded_at,
                "status": f.status,
            }

            # 🔒 Admin-only: include uploader username
            if request.user.username == "Admin" and f.user:
                item["uploaded_by"] = f.user.username

            data.append(item)
        return Response(data)


class OCRDeleteView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def delete(self, request, filename):
        try:
            ocr_file = get_object_or_404(
                OCRFile,
                file_name=filename,
                user=request.user
            )

            # Only owner or Admin can delete
            if not (request.user.username == 'Admin' or ocr_file.user == request.user):
                return Response({"error": "Forbidden"}, status=403)
            # Delete local files if they exist
            user_output_dir = Path(settings.MEDIA_ROOT) / "outputs" / str(ocr_file.user.id)
            base_name = Path(filename).stem
            
            # Try to delete all possible generated files
            for ext in ['pdf', 'txt', 'docx']:
                file_path = user_output_dir / f"ocr_{base_name}.{ext}"
                if file_path.exists():
                    file_path.unlink()
            
            # Delete the database entry
            ocr_file.delete()
            return Response({"status": "deleted"})
        except OCRFile.DoesNotExist:
            return Response({"error": "File not found"}, status=404)


class OCRDownloadView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request, filename):
        # optional query param ?ext=pdf|txt|docx
        ext = request.GET.get('ext', 'txt')
        ocr_file = get_object_or_404(
            OCRFile,
            file_name=filename,
            user=request.user
        )

        # Check permissions
        if not (request.user.username == 'Admin' or ocr_file.user == request.user):
            return Response({"error": "Forbidden"}, status=403)

        base_name = Path(filename).stem
        user_output_dir = Path(settings.MEDIA_ROOT) / "outputs" / str(ocr_file.user.id)
        file_path = user_output_dir / f"ocr_{base_name}.{ext}"

        if not file_path.exists():
            raise Http404("File not found")
        return FileResponse(open(file_path, 'rb'), as_attachment=True, filename=file_path.name)


class UserProfileView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = request.user
        # Try to get user profile, create if it doesn't exist
        try:
            profile = user.profile
            role = profile.role
        except:
            role = 'admin' if user.username == 'Admin' else 'user'
        
        data = {
            "username": user.username,
            "email": user.email,
            "is_superuser": user.is_superuser,
            "role": role,
            "first_name": user.first_name,
            "last_name": user.last_name,
        }
        return Response(data)


class AllUsersView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # Only admin can access this endpoint
        if request.user.username != 'Admin':
            return Response({"error": "Forbidden"}, status=403)
        
        # Get all users except the admin
        all_users = User.objects.all()
        users_data = []
        
        for u in all_users:
            try:
                profile = u.profile
                role = profile.role
            except:
                role = 'admin' if u.username == 'Admin' else 'user'
            
            # Get user's documents
            user_files = OCRFile.objects.filter(user=u).order_by('-uploaded_at')
            total_uploads = user_files.count()
            total_size = sum([f.pdf_size + f.txt_size + f.docx_size for f in user_files])
            
            # Get last upload date
            last_upload = None
            if total_uploads > 0:
                last_upload = user_files.first().uploaded_at.isoformat()
            
            users_data.append({
                "username": u.username,
                "email": u.email,
                "role": role,
                "total_uploads": total_uploads,
                "total_size": total_size,
                "last_upload": last_upload,
                "documents": [
                    {
                        "file_name": f.file_name,
                        "uploaded_at": f.uploaded_at,
                        "status": f.status,
                        "pdf_size": f.pdf_size,
                        "txt_size": f.txt_size,
                        "docx_size": f.docx_size
                    }
                    for f in user_files
                ]
            })
        
        return Response(users_data)


class SetSecurityQuestionsView(APIView):
    """API endpoint to set security questions for authenticated users"""
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        from .serializers import SetSecurityQuestionsSerializer
        from .models import UserProfile
        
        serializer = SetSecurityQuestionsSerializer(data=request.data)
        if serializer.is_valid():
            try:
                profile = request.user.profile
            except UserProfile.DoesNotExist:
                return Response({"error": "User profile not found"}, status=status.HTTP_404_NOT_FOUND)
            
            profile.security_answer_1 = serializer.validated_data['security_answer_1']
            profile.security_answer_2 = serializer.validated_data['security_answer_2']
            profile.save()
            
            return Response({
                "msg": "Security questions set successfully",
                "question_1": profile.security_question_1,
                "question_2": profile.security_question_2
            }, status=status.HTTP_200_OK)
        
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    def get(self, request):
        """Get security questions for authenticated user"""
        from .models import UserProfile
        
        try:
            profile = request.user.profile
        except UserProfile.DoesNotExist:
            return Response({"error": "User profile not found"}, status=status.HTTP_404_NOT_FOUND)
        
        return Response({
            "question_1": profile.security_question_1,
            "question_2": profile.security_question_2,
            "is_set": profile.security_answer_1 is not None and profile.security_answer_2 is not None
        }, status=status.HTTP_200_OK)


class ForgotPasswordView(APIView):
    """API endpoint for password reset using security questions"""
    permission_classes = [AllowAny]
    
    def post(self, request):
        from .serializers import ForgotPasswordSerializer
        
        serializer = ForgotPasswordSerializer(data=request.data)
        if serializer.is_valid():
            user = serializer.validated_data['user']
            new_password = serializer.validated_data['new_password']
            
            user.set_password(new_password)
            user.save()
            
            return Response({
                "msg": "Password reset successfully. Please login with your new password."
            }, status=status.HTTP_200_OK)
        
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class ChangePasswordView(APIView):
    """API endpoint for authenticated users to change password"""
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        from .serializers import ChangePasswordSerializer
        
        serializer = ChangePasswordSerializer(data=request.data)
        if serializer.is_valid():
            user = request.user
            old_password = serializer.validated_data['old_password']
            new_password = serializer.validated_data['new_password']
            
            # Verify old password
            if not user.check_password(old_password):
                return Response({
                    "old_password": "Old password is incorrect"
                }, status=status.HTTP_400_BAD_REQUEST)
            
            user.set_password(new_password)
            user.save()
            
            return Response({
                "msg": "Password changed successfully"
            }, status=status.HTTP_200_OK)
        
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class DeleteUserView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def delete(self, request, username):
        if request.user.username != 'Admin':
            return Response({"error": "Forbidden"}, status=403)
        
        try:
            user_to_delete = User.objects.get(username=username)
            if user_to_delete.username == 'Admin':
                return Response({"error": "Cannot delete admin"}, status=400)
            user_to_delete.delete()
            return Response({"msg": "User deleted successfully"})
        except User.DoesNotExist:
            return Response({"error": "User not found"}, status=404)


class CustomTokenObtainPairView(TokenObtainPairView):
    def post(self, request, *args, **kwargs):
        username = request.data.get('username')
        if username and not User.objects.filter(username=username).exists():
            return Response({"detail": "Your account has been deleted by the admin."}, status=401)
        return super().post(request, *args, **kwargs)
    
# ===========================
# CHAT APIs
# ===========================

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def chat_users(request):
    """
    List all users except the logged-in user
    """
    users = User.objects.exclude(id=request.user.id)
    serializer = ChatUserSerializer(users, many=True)
    return Response(serializer.data)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def get_or_create_chat_room(request):
    """
    Get or create a 1-to-1 chat room
    """
    other_user_id = request.data.get('user_id')

    if not other_user_id:
        return Response({"error": "user_id is required"}, status=400)

    try:
        other_user = User.objects.get(id=other_user_id)
    except User.DoesNotExist:
        return Response({"error": "User not found"}, status=404)

    room = ChatRoom.get_or_create_room(request.user, other_user)
    serializer = ChatRoomSerializer(room)
    return Response(serializer.data)


@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def chat_messages(request, room_id):
    room = get_object_or_404(ChatRoom, id=room_id)

    if request.method == 'GET':
        messages = (
            Message.objects
            .filter(room=room)
            .exclude(deleted_for=request.user)
            .order_by('created_at')
        )

        serializer = MessageSerializer(
            messages,
            many=True,
            context={"request": request}
        )
        return Response(serializer.data)

    # ==========================
    # POST → SEND MESSAGE
    # ==========================
    elif request.method == 'POST':
        content = request.data.get("content", "").strip()

        if not content:
            return Response(
                {"error": "Message content cannot be empty"},
                status=400
            )

        message = Message.objects.create(
            room=room,
            sender=request.user,
            content=content
        )

        serializer = MessageSerializer(
            message,
            context={"request": request}
        )

        # ✅ THIS RETURN WAS MISSING
        return Response(serializer.data, status=201)



@api_view(["DELETE"])
@permission_classes([IsAuthenticated])
def delete_message_for_everyone(request, message_id):
    message = get_object_or_404(Message, id=message_id)

    # 🔐 STRICT check
    if message.sender_id != request.user.id and not request.user.is_superuser:
        return Response({"error": "Forbidden"}, status=403)

    message.is_deleted_for_everyone = True
    message.content = "This message was deleted"
    message.save(update_fields=["is_deleted_for_everyone", "content"])

    return Response({"status": "deleted_for_everyone"})


@api_view(["PUT"])
@permission_classes([IsAuthenticated])
def edit_message(request, message_id):
    message = get_object_or_404(Message, id=message_id)

    if message.sender != request.user:
        return Response({"error": "Forbidden"}, status=403)

    if message.is_deleted_for_everyone:
        return Response({"error": "Cannot edit deleted message"}, status=400)

    new_content = request.data.get("content", "").strip()
    if not new_content:
        return Response({"error": "Message cannot be empty"}, status=400)

    message.content = new_content
    message.is_edited = True
    message.edited_at = timezone.now()
    message.save()

    return Response(MessageSerializer(message).data)

@api_view(["DELETE"])
@permission_classes([IsAuthenticated])
def delete_message_for_me(request, message_id):
    """
    Deletes a message only for the current user
    """
    message = get_object_or_404(Message, id=message_id)

    # Add current user to deleted_for list
    message.deleted_for.add(request.user)

    return Response({"status": "deleted_for_me"})
